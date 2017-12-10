import docx
from nltk.tokenize import word_tokenize
import classifier as clf
import string
import sys
import subprocess as sp
filename = sys.argv[1]

def file_check(filename):
    e, output = sp.getstatusoutput('ls | grep ' + filename + ';')
    if e != 0:
        print("No such file in the current directory.")
        return None, -1
    else:
        print("File imported succesfully...")
        return output, 1

def add_hyperlink_to_doc(complex_list, document):

    paragraph = document.add_paragraph()
    gen_url = 'https://www.merriam-webster.com/dictionary/'

    num = 1
    for c in complex_list:

        url = gen_url + c
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together and add the required text to the w:r element

        blank_space_and_num = docx.oxml.shared.OxmlElement('w:r')
        blank_space_and_num.text = '\n\t' + str(num) + '. '

        paragraph._p.append(blank_space_and_num)

        new_run.append(rPr)
        new_run.text = c
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        num += 1

def save_complex_total(input_file):

    file_as_text = open(input_file, 'r').read()
    file_as_string = file_as_text.translate(str.maketrans('', '', string.punctuation))

    all_words_list = word_tokenize(file_as_string)

    complex_list = []

    for w in all_words_list:
        word = w.strip().lower()
        if clf.classifier(word) == 1:
            if word not in complex_list:
                complex_list.append(word)

    document = docx.Document()
    para_original_text = document.add_paragraph()

    original_text = file_as_text + '\n\nThe following are hyperlinks to meanings of complex words found :- '
    text_element = docx.oxml.shared.OxmlElement('w:r')
    text_element.text = original_text
    para_original_text._p.append(text_element)

    add_hyperlink_to_doc(complex_list = complex_list, document = document)

    document.save(str(input_file.split('.')[0]) + "_tagged_words.docx")

filename, status = file_check(filename)

if status == -1:
    print("Exiting...")
    exit()
else:
    save_complex_total(filename)
    print("The doc file has been saved by the name :" + str(filename.split('.')[0]) + "_tagged_words.docx")
# punc = [ele for ele in string.punctuation]
# print(punc)
